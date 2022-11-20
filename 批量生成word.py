from win32com.client import Dispatch
from win32com.client.gencache import EnsureDispatch
from win32com.client import constants
import openpyxl
from PIL import ImageGrab, Image    #PIL图像处理库
import time
import win32com.client as win32
# excel = EnsureDispatch('Excel.Application') #启动Excel程序
# excel.Visible = True  # 可视化                              #是否显示Excel界面
# excel.DisplayAlerts = False  # 是否显示警告
wordorder = EnsureDispatch('Word.Application')  #启动word程序
wordorder.Visible = True                        #程序可见
lujing = r"D:\03 python代码\01Pywin32批量写报告\03 报告\07 2021年西宁市湟源县报告-0528.docx"  #word报告模板路径
lujing2 =  r"D:\03 python代码\01Pywin32批量写报告\03 报告"                                #word报告存储路径
lujing3 =  r"D:\03 python代码\01Pywin32批量写报告\02 Excel报表\西宁市湟源县.xlsx"            #Excel数据路径
wb = openpyxl.load_workbook(lujing3,data_only=True)        #用openxl打开Excel工作簿模板   data_only=True保证从Excel粘贴到word中的数据不含公式，而是数值
sheet1 = wb['Sheet2']                     #锁定工作表
document = wordorder.Documents.Open(lujing)     #打开word模板
范围 = document.Range()                         #选中全文
################################################################################
#下面代码处理文档中的文字
for i in range(1,24):
    a = sheet1['B{}'.format(i)].value
    范围.Find.Execute('[{}]'.format(i), False, False, False, False, False, True, 1, True, '{}'.format(a), 1)    #替换文字操作
# wb.close()
################################################################################
#下面代码用openpyxl替换表格中的数字
import time
from win32com import client
from win32com.client import constants
from win32com.client import Dispatch
import win32clipboard
sheet2 = wb['0全省PQI']
table4 = document.Tables(10)
#为单元格赋值
for i in range(1,table4.Rows.Count+1):
    for j in range(1,table4.Columns.Count+1):
        a = sheet2.cell(row=i+1,column=j+16).value
        table4.Cell(i,j).Range.Text = a             #直接Cell(i,j)无法赋值，必须转成万能的range，然后就可以为所欲为了
        # table4.Cell(i,j).Range.Font.Name =  "Forte"

#为单元格背景添加颜色

wb.close()
################################################################################
# 下面代码处理word中的图片问题
excel = win32.gencache.EnsureDispatch('Excel.Application') #启动Excel程序
excel.Visible = True  # 可视化                              #是否显示Excel界面
excel.DisplayAlerts = False  # 是否显示警告
workbook = excel.Workbooks.Open(lujing3)                  #打开Excel工作簿
for sheet in workbook.Worksheets:
    if sheet.Name == '0全省PQI':
        for i, shape in enumerate(sheet.Shapes):
            # print(sheet.Name, shape.Name)
            if shape.Name.startswith('Chart 8'):  # 'Picture'为图，Chart 为图形
                shape.Copy()
                image = ImageGrab.grabclipboard()
                image.save('temp{}.jpg'.format(i + 1), 'png')

                # print(type(image))
                # doc.add_paragraph('第{}张图'.format(i + 1))
                范围 = document.Range()
                范围.Find.Execute("（见图3-1）。")
                范围.InsertParagraphAfter()
                document.Range(范围.End, 范围.End).InlineShapes.AddPicture(
                    FileName=r"D:\03 python代码\01Pywin32批量写报告\temp2.jpg")
                # document.add_picture('temp{}.jpg'.format(i + 1), width=Inches(6.0))
################################################################################
# 下面代码处理word中的表格问题
import time
from win32com import client
from win32com.client import constants

# sheet1=workbook.Worksheets(20)
# sheet2 = workbook.Worksheets(10)
# 光标 = wordorder.Selection
# 光标.Find.Execute("[表3-2]")
# sheet2.Range('Q2:Z5').Copy()
# # document.Range(范围.End, 范围.End).Paste()
# 光标.PasteAndFormat(constants.wdFormatOriginalFormatting)#粘贴刚才复制过的表格 Paste()前面可以是Range属性，但是PastAndFormat前面必须是光标属性
# table = document.Tables(4)
# table.Columns(1).Width = 60   #调整列宽

sheet1.Range('E1:N749').Copy()# 复制表中A1到B5的范围，A1为左上角的单元格坐标，B5为右下角的坐标
parag=document.Paragraphs.Last#将变量parag指向word文档中最后一段的段尾
# parag.Range.Paste()
parag.Range.PasteAndFormat(constants.wdFormatOriginalFormatting)#粘贴刚才复制过的表格


document.SaveAs(FileName = lujing2+r"\01.docx")            #将word另存为

#为单元格背景添加颜色
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import docx
document = docx.Document(r"D:\03 python代码\01Pywin32批量写报告\03 报告\01.docx")
tb3 = document.tables[9]
for i in range(1,4):
    if float(tb3.rows[i].cells[1].text)>=90:
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="78E039"/>'.format(nsdecls('w')))
        tb3.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
    elif float(tb3.rows[i].cells[1].text)>=80:
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#61FBE7"/>'.format(nsdecls('w')))
        tb3.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
    elif float(tb3.rows[i].cells[1].text)>=70:
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="E0EE73"/>'.format(nsdecls('w')))
        tb3.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
    elif float(tb3.rows[i].cells[1].text)>=60:
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFAA52"/>'.format(nsdecls('w')))
        tb3.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
    else:
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FA5402"/>'.format(nsdecls('w')))
        tb3.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
document.save(r"D:\03 python代码\01Pywin32批量写报告\03 报告\02.docx")




