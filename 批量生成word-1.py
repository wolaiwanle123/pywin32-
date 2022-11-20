from win32com.client.gencache import EnsureDispatch
import openpyxl
from PIL import ImageGrab, Image    #PIL图像处理库
import win32com.client as win32
import os
os.chdir(r'D:\03 python代码\01Pywin32批量写报告\02 Excel报表')
filename = os.listdir()
lujing = r"D:\03 python代码\01Pywin32批量写报告\03 报告\07 2021年西宁市湟源县报告-0528.docx"  #word报告模板路径
lujing2 =  r"D:\03 python代码\01Pywin32批量写报告\03 报告"                                #word报告存储路径
lujing3 =  r"D:\03 python代码\01Pywin32批量写报告\02 Excel报表"            #Excel数据路径
lujing4 = r'D:\03 python代码\01Pywin32批量写报告\04 报告'
for k in range(0,2):
    word = EnsureDispatch('Word.Application')  #启动word程序
    word.Visible = True
    document = word.Documents.Open(lujing)     #打开word模板
    范围 = document.Range()
    wb = openpyxl.load_workbook(lujing3+'\\'+filename[k],data_only=True)        #用openxl打开Excel工作簿模板   data_only=True保证从Excel粘贴到word中的数据不含公式，而是数值
    sheet1 = wb['Sheet2']                     #锁定工作表
################################################################################
    #下面代码处理文档中的文字
    for i in range(1,24):
        a = sheet1['B{}'.format(i)].value
        范围.Find.Execute('[{}]'.format(i), False, False, False, False, False, True, 1, True, '{}'.format(a), 1)    #替换文字操作
################################################################################
    #下面代码用openpyxl替换表格中的数字
    sheet2 = wb['0全省PQI']
    #表3-1
    table4 = document.Tables(10)
    for i in range(1,table4.Rows.Count+1):
        for j in range(1,table4.Columns.Count+1):
            a = sheet2.cell(row=i+1,column=j+16).value
            table4.Cell(i,j).Range.Text = a             #直接Cell(i,j)无法赋值，必须转成万能的range，然后就可以为所欲为了
    #表3-2
    table4 = document.Tables(11)
    for i in range(1,table4.Rows.Count+1):
        for j in range(1,table4.Columns.Count+1):
            a = sheet2.cell(row=i+53,column=j+16).value
            table4.Cell(i,j).Range.Text = a
    wb.close()
################################################################################
    # 下面代码处理word中的图片问题
    excel = win32.gencache.EnsureDispatch('Excel.Application') #启动Excel程序
    excel.Visible = True  # 可视化                              #是否显示Excel界面
    excel.DisplayAlerts = False  # 是否显示警告
    workbook = excel.Workbooks.Open(lujing3+'\\'+filename[k])                  #打开Excel工作簿
    for sheet in workbook.Worksheets:
        if sheet.Name == '0全省PQI':
            for i, shape in enumerate(sheet.Shapes):
                if shape.Name.startswith('Chart 8'):  # 'Picture'为图，Chart 为图形
                    shape.Copy()
                    image = ImageGrab.grabclipboard()
                    image.save('temp{}.jpg'.format(i + 1), 'png')

                    # print(type(image))
                    # doc.add_paragraph('第{}张图'.format(i + 1))
                    范围 = document.Range()
                    范围.Find.Execute("（见图3-1）。")
                    范围.InsertParagraphAfter()
                    document.Range(范围.End, 范围.End).InlineShapes.AddPicture(
                        FileName=r"D:\03 python代码\01Pywin32批量写报告\02 Excel报表\temp2.jpg")
################################################################################
    # 下面代码处理word中的表格问题：直接将最后的公里明细附表从Excel中粘贴过来
    from win32com.client import constants
    sheet1 = workbook.Worksheets(20)
    info = sheet1.UsedRange           #查找sheet1工作表中有数据范围
    nrows = info.Rows.Count           #查找sheet1工作表中有数据行数范围
    ncols = info.Columns.Count        #查找sheet1工作表中有数据列数范围
    print(nrows)
    sheet1.Range('A1:J{}'.format(nrows)).Copy()# 复制表中A1到B5的范围，A1为左上角的单元格坐标，B5为右下角的坐标
    # print(document)
    parag=document.Paragraphs.Last#将变量parag指向word文档中最后一段的段尾
    # parag.Range.Paste()
    # document.Range(document.Range().End - 1).Paste()
    parag.Range.PasteAndFormat(constants.wdFormatOriginalFormatting)#粘贴刚才复制过的表格
################################################################################

    document.SaveAs(FileName = lujing2+'\\'+filename[k])            #将word另存为
    document.Close(SaveChanges = 0)
    workbook.Close(SaveChanges = 0)
    word.Quit()
    excel.Quit()

    # 为单元格背景添加颜色
    # talecolor()根据单元格数值赋予颜色
    # #为单元格背景添加颜色
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    import docx
    document = docx.Document(lujing2+'\\'+filename[k])
    def tablecolor(tb):
        for i in range(1, 4):
            if float(tb.rows[i].cells[1].text) >= 90:
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="78E039"/>'.format(nsdecls('w')))
                tb.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
            elif float(tb.rows[i].cells[1].text) >= 80:
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#61FBE7"/>'.format(nsdecls('w')))
                tb.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
            elif float(tb.rows[i].cells[1].text) >= 70:
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="E0EE73"/>'.format(nsdecls('w')))
                tb.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
            elif float(tb.rows[i].cells[1].text) >= 60:
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFAA52"/>'.format(nsdecls('w')))
                tb.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
            else:
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FA5402"/>'.format(nsdecls('w')))
                tb.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)


    for i in range(9, 11):
        tb = document.tables[i]
        tablecolor(tb)
    document.save(lujing4+'\\'+filename[k])
################################################################################
# #为单元格背景添加颜色
# from docx.oxml.ns import nsdecls
# from docx.oxml import parse_xml
# import docx
# document = docx.Document(r"D:\03 python代码\01Pywin32批量写报告\03 报告\01.docx")
# #talecolor()根据单元格数值赋予颜色
# def tablecolor(tb):
#     for i in range(1,4):
#         if float(tb.rows[i].cells[1].text)>=90:
#             shading_elm_1 = parse_xml(r'<w:shd {} w:fill="78E039"/>'.format(nsdecls('w')))
#             tb.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
#         elif float(tb.rows[i].cells[1].text)>=80:
#             shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#61FBE7"/>'.format(nsdecls('w')))
#             tb.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
#         elif float(tb.rows[i].cells[1].text)>=70:
#             shading_elm_1 = parse_xml(r'<w:shd {} w:fill="E0EE73"/>'.format(nsdecls('w')))
#             tb.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
#         elif float(tb.rows[i].cells[1].text)>=60:
#             shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFAA52"/>'.format(nsdecls('w')))
#             tb.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
#         else:
#             shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FA5402"/>'.format(nsdecls('w')))
#             tb.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
# for i in range(9,11):
#     tb = document.tables[i]
#     tablecolor(tb)
# document.save(r"D:\03 python代码\01Pywin32批量写报告\03 报告\02.docx")